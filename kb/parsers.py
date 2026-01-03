from __future__ import annotations

import csv
import io
import os
from typing import Optional

from PIL import Image

try:
    from docx import Document  # python-docx
except Exception:
    Document = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    import openpyxl
except Exception:
    openpyxl = None


def detect_ext(path: str) -> str:
    return (os.path.splitext(path)[1] or "").lower().strip(".")


def parse_text_bytes(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1251"):
        try:
            return data.decode(enc)
        except Exception:
            continue
    return data.decode("utf-8", errors="ignore")


def parse_pdf_bytes(data: bytes) -> str:
    if PdfReader is None:
        raise RuntimeError("pypdf is not installed")
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n".join(parts).strip()


def parse_docx_bytes(data: bytes) -> str:
    if Document is None:
        raise RuntimeError("python-docx is not installed")
    doc = Document(io.BytesIO(data))
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # таблицы из DOCX
    for t in doc.tables:
        for row in t.rows:
            parts.append("\t".join([c.text or "" for c in row.cells]))
    return "\n".join(parts).strip()


def parse_csv_bytes(data: bytes) -> str:
    text = parse_text_bytes(data)
    buf = io.StringIO(text)
    reader = csv.reader(buf)
    lines = []
    for row in reader:
        lines.append("\t".join(row))
    return "\n".join(lines).strip()


def parse_xlsx_bytes(data: bytes) -> str:
    if openpyxl is None:
        raise RuntimeError("openpyxl is not installed")
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    out = []
    for sheet in wb.worksheets:
        out.append(f"[SHEET] {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            if row is None:
                continue
            vals = []
            for v in row:
                if v is None:
                    vals.append("")
                else:
                    vals.append(str(v))
            line = "\t".join(vals).strip()
            if line:
                out.append(line)
    return "\n".join(out).strip()


def is_image_ext(ext: str) -> bool:
    return ext in {"png", "jpg", "jpeg", "webp"}


def parse_image_bytes_best_effort(data: bytes) -> str:
    """
    Без OCR — только метаданные/placeholder.
    OCR/vision лучше подключать опционально через OpenAI Responses API (см. indexer/syncer).
    """
    try:
        im = Image.open(io.BytesIO(data))
        return f"[IMAGE] format={im.format} size={im.size[0]}x{im.size[1]}"
    except Exception:
        return "[IMAGE] unable_to_parse"
